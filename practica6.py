# Librerías utilizadas para el proyecto

from flask import Flask, render_template, send_from_directory, request
import os
from PyPDF2 import PdfReader
from docx import Document
import pandas as pd
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.stem import SnowballStemmer
import string

# Recursos necesarios de NLTK
nltk.download("punkt")
nltk.download("stopwords")
nltk.download("punkt_tab")


# Iniciación de Flask para la GUI
app = Flask(__name__)

# Ruta del corpus a utilizar
CORPUS_FOLDER = r'D:\Documentos\Luis\CCO\Recuperacion Informacion\RepositorioParaPractica3RI'

#  FUNCIONES PARA CONVERSIÓN BOOLEANA

# Prioridad de operadores
prioridad = {
    "no": 3,
    "y": 2,
    "o": 1
}

def limpiar_token(token):
    import re
    import unicodedata

    # Normalizar Unicode (elimina raro “\u200b”)
    token = unicodedata.normalize("NFKD", token).replace("\u200b", "")

    # Eliminar símbolos al inicio o fin
    token = re.sub(r'^[¿¡“”"«»•—–\-]+', '', token)
    token = re.sub(r'[¿¡“”"«»•—–\-]+$', '', token)

    # Deja solo letras (incluye tildes y ñ)
    token = re.sub(r'[^a-zA-ZñÑáéíóúÁÉÍÓÚ]', '', token)

    token = token.lower()

    # Eliminar palabras de 1 letra o vacías
    if len(token) < 2:
        return ""

    return token


def es_operador(token):
    return token in ["y", "o", "no"]

def a_posfijo(tokens):
    """Convierte expresión infijo a posfijo"""
    salida = []
    pila = []

    for token in tokens:
        if token == "(":
            pila.append(token)
        elif token == ")":
            while pila and pila[-1] != "(":
                salida.append(pila.pop())
            pila.pop()  # quitar "("
        elif es_operador(token):
            while (pila and pila[-1] != "(" and 
                   prioridad.get(pila[-1], 0) >= prioridad[token]):
                salida.append(pila.pop())
            pila.append(token)
        else:
            salida.append(token)

    while pila:
        salida.append(pila.pop())

    return salida


def evaluar_posfijo(tokens):
    """Lectura paso a paso del posfijo"""
    pila = []
    print("\n--- LECTURA POSFIJO PASO A PASO ---\n")
    for token in tokens:
        print(f"Procesando: {token}")
        if not es_operador(token):
            pila.append(token)
            print(f"  Apila → {pila}")
        else:
            if token == "no":
                op = pila.pop()
                nuevo = f"(no {op})"
                pila.append(nuevo)
            else:
                right = pila.pop()
                left = pila.pop()
                nuevo = f"({left} {token} {right})"
                pila.append(nuevo)
            
            print(f"  Resultado parcial: {pila[-1]}")
        
    print("\nResultado final:", pila[0])
    print("-------------------------------------\n")

# Definición del nodo para el árbol B
class NodoB:
    def __init__(self, hoja=False):
        self.hoja = hoja
        self.claves = []
        self.hijos = []

# Creación de la clase para el árbol B
class ArbolB:
    def __init__(self, grado=4):
        self.raiz = NodoB(hoja=True)
        self.grado = grado
        self.max_claves = grado - 1

# Función para insertar nodos
    def insertar(self, clave):
        r = self.raiz
        # Si nodo lleno, dividir
        if len(r.claves) == self.max_claves:
            nuevo = NodoB()
            self.raiz = nuevo
            nuevo.hijos.append(r)
            self._split(nuevo, 0)
            self._insert_non_full(nuevo, clave)
        else:
            self._insert_non_full(r, clave)

    def _insert_non_full(self, nodo, clave):
        if nodo.hoja:
            nodo.claves.append(clave)
            nodo.claves.sort()
        else:
            i = len(nodo.claves) - 1
            while i >= 0 and clave < nodo.claves[i]:
                i -= 1
            i += 1

            if len(nodo.hijos[i].claves) == self.max_claves:
                self._split(nodo, i)
                if clave > nodo.claves[i]:
                    i += 1
            self._insert_non_full(nodo.hijos[i], clave)

    def _split(self, nodo_padre, index):
        nodo = nodo_padre.hijos[index]
        medio = self.max_claves // 2
        clave_medio = nodo.claves[medio]

        nuevo = NodoB(hoja=nodo.hoja)
        nuevo.claves = nodo.claves[medio+1:]
        nodo.claves = nodo.claves[:medio]

        if not nodo.hoja:
            nuevo.hijos = nodo.hijos[medio+1:]
            nodo.hijos = nodo.hijos[:medio+1]

        nodo_padre.claves.insert(index, clave_medio)
        nodo_padre.hijos.insert(index+1, nuevo)

    def buscar(self, nodo, clave):
        i = 0
        while i < len(nodo.claves) and clave > nodo.claves[i]:
            i += 1

        if i < len(nodo.claves) and clave == nodo.claves[i]:
            return True

        if nodo.hoja:
            return False
        else:
            return self.buscar(nodo.hijos[i], clave)

    def imprimir(self):
        print("\n=== ÁRBOL B ===")
        self.raiz_imprimir(self.raiz)
        print("================\n")


    def raiz_imprimir(self, nodo, nivel=0):
        print("  " * nivel + str(nodo.claves))
        if not nodo.hoja:
            for hijo in nodo.hijos:
                self.raiz_imprimir(hijo, nivel+1)

#Funciones para leer cualquier documento txt, doc, o pdf
def leer_documento_completo(ruta_archivo):
    archivo = ruta_archivo.lower()

    if archivo.endswith(".txt"):
        with open(ruta_archivo, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    elif archivo.endswith(".pdf"):
        texto = ""
        reader = PdfReader(ruta_archivo)
        for page in reader.pages:
            contenido = page.extract_text()
            if contenido:
                texto += contenido + "\n"
        return texto

    elif archivo.endswith(".docx"):
        texto = ""
        doc = Document(ruta_archivo)
        for p in doc.paragraphs:
            texto += p.text + "\n"
        return texto

    return ""

def ejecutar_preprocesamiento():
    print("\n===== PREPROCESAMIENTO INICIADO =====\n")

    diccionario = set()

    # -------------------
    # TOKENIZACIÓN
    # -------------------
    for archivo in os.listdir(CORPUS_FOLDER):
        ruta = os.path.join(CORPUS_FOLDER, archivo)

        if not os.path.isfile(ruta):
            continue

        texto = leer_documento_completo(ruta)

        if texto.strip():
            tokens = word_tokenize(texto)

            # 🔥 Aplicar limpieza desde el principio
            tokens = [limpiar_token(t) for t in tokens]
            tokens = [t for t in tokens if t]  # quitar vacíos

            diccionario.update(tokens)

    # Guardar diccionario original ya limpio
    df_tokens = pd.DataFrame(sorted(diccionario), columns=["Token"])
    df_tokens.to_excel("1Diccionario.xlsx", index=False)
    print("Archivo generado: 1Diccionario.xlsx")

    # -------------------
    # LIMPIEZA PROFUNDA (segunda pasada)
    # -------------------
    tokens_limpios = []

    for tok in diccionario:
        limpio = limpiar_token(tok)
        if limpio:
            tokens_limpios.append(limpio)

    df_limpio = pd.DataFrame(sorted(set(tokens_limpios)), columns=["Token"])
    df_limpio.to_excel("2DiccMinus.xlsx", index=False)
    print("Archivo generado: 2DiccMinus.xlsx")

    # -------------------
    # ELIMINAR STOPWORDS
    # -------------------
    stop_words = set(stopwords.words("spanish"))
    sin_stopwords = []

    for p in df_limpio["Token"]:
        p_clean = limpiar_token(p)  # volver a limpiar antes de revisar stopwords, en caso de haber faltado algo
        if p_clean and p_clean not in stop_words:
            sin_stopwords.append(p_clean)

    df_stop = pd.DataFrame(sorted(set(sin_stopwords)), columns=["Token"])
    df_stop.to_excel("3DiccSinStopWords.xlsx", index=False)
    print("Archivo generado: 3DiccSinStopWords.xlsx")

    # -------------------
    # STEMMING
    # -------------------
    stemmer = SnowballStemmer("spanish")
    stems = []

    for p in df_stop["Token"]:
        base = limpiar_token(p) 
        if base:
            stems.append(stemmer.stem(base))

    df_stem = pd.DataFrame(sorted(set(stems)), columns=["Token"])
    df_stem.to_excel("4DiccSteams.xlsx", index=False)
    print("Archivo generado: 4DiccSteams.xlsx")

    print("\n===== PREPROCESAMIENTO COMPLETO =====\n")

def generar_matriz_binaria():
    print("\n===== GENERANDO MATRIZ BINARIA DE STEAMS =====\n")

    df_stems = pd.read_excel("4DiccSteams.xlsx")
    vocabulario = sorted(df_stems["Token"].tolist())

    matriz = []

    for archivo in os.listdir(CORPUS_FOLDER):
        ruta = os.path.join(CORPUS_FOLDER, archivo)

        if not os.path.isfile(ruta):
            continue

        texto = leer_documento_completo(ruta)

        tokens = word_tokenize(texto.lower())

        tokens = [
            t.translate(str.maketrans("", "", string.punctuation))
            for t in tokens
        ]
        tokens = [t for t in tokens if t.strip()]

        stop_words = set(stopwords.words("spanish"))
        tokens = [t for t in tokens if t not in stop_words]

        stemmer = SnowballStemmer("spanish")
        steams_doc = [stemmer.stem(t) for t in tokens]

        fila = [1 if steam in steams_doc else 0 for steam in vocabulario]

        matriz.append([archivo] + fila)

    columnas = ["Documento"] + vocabulario
    df_matriz = pd.DataFrame(matriz, columns=columnas)

    df_matriz.to_excel("6MatrizBinariaDeSteams.xlsx", index=False)

    print("Matriz generada: 6MatrizBinariaDeSteams.xlsx\n")
    print("===== MATRIZ COMPLETA LISTA =====\n")

def procesar_consulta_booleana_y_vector(expresion):
    global arbol_consulta, tabla_hash_consulta

    print("\n===== PROCESANDO CONSULTA: VECTOR + HASH + ÁRBOL B =====\n")
    print("Consulta original:", expresion)

    # ------------------------------------------------------------------
    # PREPROCESAR CONSULTA
    # ------------------------------------------------------------------
    expresion = expresion.lower()
    expresion_clean = expresion.replace("(", " ( ").replace(")", " ) ")
    tokens = word_tokenize(expresion_clean)

    tokens = [
        t.translate(str.maketrans("", "", string.punctuation))
        for t in tokens
    ]

    tokens = [t for t in tokens if t.strip()]

    stop_words = set(stopwords.words("spanish"))
    tokens = [t for t in tokens if t not in stop_words]

    stemmer = SnowballStemmer("spanish")
    steams_Q = [
        stemmer.stem(t)
        for t in tokens
        if t not in ["y", "o", "no"]
    ]

    print("Steams de la consulta:", steams_Q)

    # LEER VOCABULARIO
    df_vocab = pd.read_excel("4DiccSteams.xlsx")
    vocabulario = df_vocab["Token"].tolist()

    print("\nVocabulario del corpus:")
    print(vocabulario)

    
    # VECTOR BINARIO DE CONSULTA
    
    vector = [1 if v in steams_Q else 0 for v in vocabulario]

    print("\nVECTOR DE CONSULTA Q:")
    print("Q:", " ".join(str(x) for x in vector))

    # TABLA HASH

    print("\n=== INSERTANDO EN TABLA HASH ===")

    tabla_hash_consulta = {}
    for steam in steams_Q:
        tabla_hash_consulta[steam] = hash(steam)
        print(f"{steam} → {tabla_hash_consulta[steam]}")


    # ÁRBOL B

    print("\n=== INSERTANDO EN ÁRBOL B ===")

    arbol_consulta = ArbolB()
    for steam in steams_Q:
        arbol_consulta.insertar(steam)

    arbol_consulta.imprimir()


    # BUSCAR DOCUMENTOS QUE CONTENGAN LOS STEAMS

    print("\n=== BUSCANDO DOCUMENTOS ===")

    df_matriz = pd.read_excel("6MatrizBinariaDeSteams.xlsx")

    documentos_encontrados = []

    for i, fila in df_matriz.iterrows():

        documento = fila.iloc[0]  # nombre del documento
        valores = fila.iloc[1:].tolist()

        coincide = False

        for j, v in enumerate(vector):
            if v == 1 and valores[j] == 1:
                coincide = True
                break

        if coincide:
            documentos_encontrados.append(documento)

    print("\nDocumentos encontrados:")
    print(documentos_encontrados)

    print("\n===== PROCESO COMPLETO =====\n")

    # Devolución de los documentos, para mostrarlos en pantalla
    return documentos_encontrados


#  RUTAS DE FLASK --------------------------------------------------------------------------------------------

@app.route("/", methods=["GET", "POST"])
def index():
    archivos = os.listdir(CORPUS_FOLDER)

    if request.method == "POST":
        print(">> Ejecutando preprocesamiento completo...")
        ejecutar_preprocesamiento()
        generar_matriz_binaria()

    return render_template("index.html", archivos=archivos)

@app.route("/documento/<nombre>")
def abrir_documento(nombre):
    ruta = os.path.join(CORPUS_FOLDER, nombre)

    if nombre.endswith(".txt"):
        with open(ruta, "r", encoding="utf-8") as f:
            print("\n--- CONTENIDO TXT ---\n")
            print(f.read())

    elif nombre.endswith(".pdf"):
        reader = PdfReader(ruta)
        print("\n--- CONTENIDO PDF ---\n")
        for pagina in reader.pages:
            print(pagina.extract_text())

    elif nombre.endswith(".docx"):
        doc = Document(ruta)
        print("\n--- CONTENIDO DOCX ---\n")
        for p in doc.paragraphs:
            print(p.text)

    return send_from_directory(CORPUS_FOLDER, nombre)


@app.route("/consulta_booleana", methods=["POST"])
def consulta_booleana():

    expresion = request.form["expresion"]

    documentos_resultado = procesar_consulta_booleana_y_vector(expresion)

    print("\n\n=====================================")
    print("EXPRESIÓN ORIGINAL:", expresion)
    print("=====================================")

    # Tokenizar separando paréntesis
    expresion_clean = (
        expresion.replace("(", " ( ")
                 .replace(")", " ) ")
    )

    tokens = expresion_clean.split()

    print("\nTokens:", tokens)

    # Convertir a posfijo
    posfijo = a_posfijo(tokens)

    print("\nNotación posfijo:", posfijo)

    # Evaluación paso a paso
    evaluar_posfijo(posfijo)
    
    print("\nDocumentos encontrados:")
    print(documentos_resultado)


    # Mostrar resultados en el navegador
    return render_template(
        "index.html",
        consulta=expresion,
        documentos=documentos_resultado
    )

from openpyxl import Workbook

tabla_hash = {}  # Diccionario global para mantener el índice
arbol_consulta = ArbolB()
tabla_hash_consulta = {}

def obtener_steams_de_archivo(ruta):
    """Extrae palabras del archivo (steams simples)."""
    texto = ""

    if ruta.endswith(".txt"):
        with open(ruta, "r", encoding="utf-8") as f:
            texto = f.read()

    elif ruta.endswith(".pdf"):
        reader = PdfReader(ruta)
        for p in reader.pages:
            texto += p.extract_text() + " "

    elif ruta.endswith(".docx"):
        doc = Document(ruta)
        for p in doc.paragraphs:
            texto += p.text + " "

    # limpieza simple
    texto = texto.lower()
    reemplazos = ",.;:()!?¡¿\n\t"
    for c in reemplazos:
        texto = texto.replace(c, " ")

    palabras = texto.split()
    return palabras

def consulta_booleana():
    expresion = request.form["expresion"]

    print("\nCONSULTA BOOLEANA \n")
    print("Expresión original:", expresion)

    expresion_clean = expresion.replace("(", " ( ").replace(")", " ) ")
    tokens = expresion_clean.split()

    print("\nTokens:", tokens)

    posfijo = a_posfijo(tokens)
    print("\nNotación posfijo:", posfijo)

    evaluar_posfijo(posfijo)

    procesar_consulta_booleana_y_vector(expresion)

    return "<h2>Consulta procesada. Revisa la consola.</h2>"

@app.route("/indexar_hash", methods=["POST"])
def indexar_hash():
    print("\n====== INDEXACIÓN HASH USANDO STEAMS ======\n")

    try:
        # 1. Cargar archivo 4 (Steams finales)
        df_stems = pd.read_excel("4DiccSteams.xlsx")

        print(f"Total de steams encontrados: {len(df_stems)}\n")

        # 2. Aplicar función hash a cada steam
        tabla_hash = {}

        for steam in df_stems["Token"]:
            valor_hash = hash(steam)
            tabla_hash[steam] = valor_hash

        # 3. Guardar archivo 5
        df_hash = pd.DataFrame(list(tabla_hash.items()), columns=["Steam", "Hash"])
        df_hash.to_excel("5ListDiccIndex.xlsx", index=False)

        print("Archivo generado: 5ListDiccIndex.xlsx\n")

        # 4. PROCESAR BÚSQUEDA QUE PIDA EL USUARIO
        steam_buscar = request.form["steam_buscar"]
        steam_buscar = steam_buscar.lower()

        print(f"\nBuscando steam: {steam_buscar}")

        if steam_buscar in tabla_hash:
            print("🔎 Resultado: ENCONTRADO")
            print(f"   Hash: {tabla_hash[steam_buscar]}\n")
        else:
            print("❌ Resultado: NO ENCONTRADO\n")

        return "<h3>Indexación Hash realizada. Revisa la consola.</h3>"

    except Exception as e:
        print("\n❌ ERROR en indexación hash:", e, "\n")
        return "<h3>Error en la indexación. Revisa consola.</h3>"
    
    
@app.route("/procesar_arbol_b", methods=["POST"])
def procesar_arbol_b():
    palabra_query = request.form["palabra_b"].lower().strip()

    # Leer archivo Excel generado anteriormente
    import openpyxl
    wb = openpyxl.load_workbook("5ListDiccIndex.xlsx")
    ws = wb.active

    print("\n===== CREANDO ÁRBOL B =====")
    arbol = ArbolB(grado=4)

    steams = []

    for row in ws.iter_rows(min_row=2, values_only=True):
        steam = row[0]
        steams.append(steam)
        arbol.insertar(steam)

    print("\nÁRBOL B COMPLETO:")
    arbol.imprimir(arbol.raiz)

    print("\n===== BÚSQUEDA EN ÁRBOL B =====")
    print(f"Buscando: {palabra_query}")

    if arbol.buscar(arbol.raiz, palabra_query):
        print("Resultado: ENCONTRADO")
        return "<h3>Palabra ENCONTRADA en el árbol B.</h3>"
    else:
        print("Resultado: NO ENCONTRADA")
        return "<h3>Palabra NO encontrada en el árbol B.</h3>"

if __name__ == "__main__":
    app.run(debug=True)