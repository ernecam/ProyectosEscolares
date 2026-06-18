from flask import Flask, render_template, send_from_directory, request
import os
from PyPDF2 import PdfReader
from docx import Document
import pandas as pd
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.stem import SnowballStemmer
import string

nltk.download("punkt")
nltk.download("stopwords")
nltk.download("punkt_tab")



app = Flask(__name__)

# CORPUS_FOLDER = ruta a corpus

# ============================================================
#  FUNCIONES PARA CONVERSIÓN BOOLEANA
# ============================================================

# Prioridad de operadores
prioridad = {
    "no": 3,
    "y": 2,
    "o": 1
}

def es_operador(token):
    return token in ["y", "o", "no"]

def a_posfijo(tokens):
    """Convierte expresión infijo a posfijo (Polaca inversa)"""
    salida = []
    pila = []

    for token in tokens:
        if token == "(":
            pila.append(token)
        elif token == ")":
            while pila and pila[-1] != "(":
                salida.append(pila.pop())
            pila.pop()  # quitar "("
        elif es_operador(token):
            while (pila and pila[-1] != "(" and 
                   prioridad.get(pila[-1], 0) >= prioridad[token]):
                salida.append(pila.pop())
            pila.append(token)
        else:
            salida.append(token)

    while pila:
        salida.append(pila.pop())

    return salida


def evaluar_posfijo(tokens):
    """Simula la lectura paso a paso del posfijo"""
    pila = []
    print("\n--- LECTURA POSFIJO PASO A PASO ---\n")
    for token in tokens:
        print(f"Procesando: {token}")
        if not es_operador(token):
            pila.append(token)
            print(f"  Apila → {pila}")
        else:
            if token == "no":
                op = pila.pop()
                nuevo = f"(no {op})"
                pila.append(nuevo)
            else:
                right = pila.pop()
                left = pila.pop()
                nuevo = f"({left} {token} {right})"
                pila.append(nuevo)
            
            print(f"  Resultado parcial: {pila[-1]}")
        
    print("\nResultado final:", pila[0])
    print("-------------------------------------\n")

class NodoB:
    def __init__(self, hoja=False):
        self.hoja = hoja
        self.claves = []
        self.hijos = []

class ArbolB:
    def __init__(self, grado=4):
        self.raiz = NodoB(hoja=True)
        self.grado = grado
        self.max_claves = grado - 1

    def insertar(self, clave):
        r = self.raiz
        # Si nodo lleno, dividir
        if len(r.claves) == self.max_claves:
            nuevo = NodoB()
            self.raiz = nuevo
            nuevo.hijos.append(r)
            self._split(nuevo, 0)
            self._insert_non_full(nuevo, clave)
        else:
            self._insert_non_full(r, clave)

    def _insert_non_full(self, nodo, clave):
        if nodo.hoja:
            nodo.claves.append(clave)
            nodo.claves.sort()
        else:
            i = len(nodo.claves) - 1
            while i >= 0 and clave < nodo.claves[i]:
                i -= 1
            i += 1

            if len(nodo.hijos[i].claves) == self.max_claves:
                self._split(nodo, i)
                if clave > nodo.claves[i]:
                    i += 1
            self._insert_non_full(nodo.hijos[i], clave)

    def _split(self, nodo_padre, index):
        nodo = nodo_padre.hijos[index]
        medio = self.max_claves // 2
        clave_medio = nodo.claves[medio]

        nuevo = NodoB(hoja=nodo.hoja)
        nuevo.claves = nodo.claves[medio+1:]
        nodo.claves = nodo.claves[:medio]

        if not nodo.hoja:
            nuevo.hijos = nodo.hijos[medio+1:]
            nodo.hijos = nodo.hijos[:medio+1]

        nodo_padre.claves.insert(index, clave_medio)
        nodo_padre.hijos.insert(index+1, nuevo)

    def buscar(self, nodo, clave):
        i = 0
        while i < len(nodo.claves) and clave > nodo.claves[i]:
            i += 1

        if i < len(nodo.claves) and clave == nodo.claves[i]:
            return True

        if nodo.hoja:
            return False
        else:
            return self.buscar(nodo.hijos[i], clave)

    def imprimir(self, nodo, nivel=0):
        print("  " * nivel + str(nodo.claves))
        if not nodo.hoja:
            for h in nodo.hijos:
                self.imprimir(h, nivel+1)

def leer_documento_completo(ruta_archivo):
    archivo = ruta_archivo.lower()

    # Leer TXT
    if archivo.endswith(".txt"):
        with open(ruta_archivo, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    # Leer PDF
    elif archivo.endswith(".pdf"):
        texto = ""
        reader = PdfReader(ruta_archivo)
        for page in reader.pages:
            contenido = page.extract_text()
            if contenido:
                texto += contenido + "\n"
        return texto

    # Leer DOCX
    elif archivo.endswith(".docx"):
        texto = ""
        doc = Document(ruta_archivo)
        for p in doc.paragraphs:
            texto += p.text + "\n"
        return texto

    return ""

def ejecutar_preprocesamiento():
    print("\n===== PREPROCESAMIENTO INICIADO =====\n")

    diccionario = set()

    # 1. TOKENIZACIÓN
    for archivo in os.listdir(CORPUS_FOLDER):
        ruta = os.path.join(CORPUS_FOLDER, archivo)

        if not os.path.isfile(ruta):
            continue

        texto = leer_documento_completo(ruta)

        if texto.strip():
            tokens = word_tokenize(texto)
            diccionario.update(tokens)

    df_tokens = pd.DataFrame(sorted(diccionario), columns=["Token"])
    df_tokens.to_excel("1Diccionario.xlsx", index=False)
    print("✔ Archivo generado: 1Diccionario.xlsx")

    # 2. MINÚSCULAS + LIMPIEZA
    diccionario_limpio = []

    for palabra in df_tokens["Token"]:
        palabra = palabra.lower()
        palabra = palabra.translate(str.maketrans("", "", string.punctuation))
        if palabra.strip():
            diccionario_limpio.append(palabra)

    df_limpio = pd.DataFrame(sorted(set(diccionario_limpio)), columns=["Token"])
    df_limpio.to_excel("2DiccMinus.xlsx", index=False)
    print("✔ Archivo generado: 2DiccMinus.xlsx")

    # 3. STOPWORDS
    stop_words = set(stopwords.words("spanish"))
    sin_stopwords = [p for p in df_limpio["Token"] if p not in stop_words]

    df_stop = pd.DataFrame(sorted(set(sin_stopwords)), columns=["Token"])
    df_stop.to_excel("3DiccSinStopWords.xlsx", index=False)
    print("✔ Archivo generado: 3DiccSinStopWords.xlsx")

    # 4. STEMMING
    stemmer = SnowballStemmer("spanish")
    stems = [stemmer.stem(p) for p in df_stop["Token"]]

    df_stem = pd.DataFrame(sorted(set(stems)), columns=["Token"])
    df_stem.to_excel("4DiccSteams.xlsx", index=False)
    print("✔ Archivo generado: 4DiccSteams.xlsx")

    print("\n🎉 PREPROCESAMIENTO COMPLETO.\n")



# ============================================================
#  RUTAS DE FLASK
# ============================================================

@app.route("/", methods=["GET", "POST"])
def index():
    archivos = os.listdir(CORPUS_FOLDER)

    if request.method == "POST":
        print(">> Ejecutando preprocesamiento completo...")
        ejecutar_preprocesamiento()

    return render_template("index.html", archivos=archivos)

@app.route("/documento/<nombre>")
def abrir_documento(nombre):
    ruta = os.path.join(CORPUS_FOLDER, nombre)

    if nombre.endswith(".txt"):
        with open(ruta, "r", encoding="utf-8") as f:
            print("\n--- CONTENIDO TXT ---\n")
            print(f.read())

    elif nombre.endswith(".pdf"):
        reader = PdfReader(ruta)
        print("\n--- CONTENIDO PDF ---\n")
        for pagina in reader.pages:
            print(pagina.extract_text())

    elif nombre.endswith(".docx"):
        doc = Document(ruta)
        print("\n--- CONTENIDO DOCX ---\n")
        for p in doc.paragraphs:
            print(p.text)

    return send_from_directory(CORPUS_FOLDER, nombre)



@app.route("/consulta_booleana", methods=["POST"])
def consulta_booleana():
    expresion = request.form["expresion"]

    print("\n\n=====================================")
    print("EXPRESIÓN ORIGINAL:", expresion)
    print("=====================================")

    # Tokenizar separando paréntesis y palabras
    expresion_clean = (
        expresion.replace("(", " ( ")
                 .replace(")", " ) ")
    )
    tokens = expresion_clean.split()

    print("\nTokens:", tokens)

    # Convertir a posfijo
    posfijo = a_posfijo(tokens)

    print("\nNotación posfijo:", posfijo)

    # Evaluar posfijo paso a paso
    evaluar_posfijo(posfijo)

    return "<h2>Consulta procesada. Revisa la consola.</h2>"

from openpyxl import Workbook

tabla_hash = {}  # Diccionario global para mantener el índice


def obtener_steams_de_archivo(ruta):
    """Extrae palabras del archivo (steams simples)."""
    texto = ""

    if ruta.endswith(".txt"):
        with open(ruta, "r", encoding="utf-8") as f:
            texto = f.read()

    elif ruta.endswith(".pdf"):
        reader = PdfReader(ruta)
        for p in reader.pages:
            texto += p.extract_text() + " "

    elif ruta.endswith(".docx"):
        doc = Document(ruta)
        for p in doc.paragraphs:
            texto += p.text + " "

    # limpieza simple
    texto = texto.lower()
    reemplazos = ",.;:()!?¡¿\n\t"
    for c in reemplazos:
        texto = texto.replace(c, " ")

    palabras = texto.split()
    return palabras



@app.route("/indexar_hash", methods=["POST"])
def indexar_hash():
    print("\n====== INDEXACIÓN HASH USANDO STEAMS ======\n")

    try:
        # 1. Cargar archivo 4 (Steams finales)
        df_stems = pd.read_excel("4DiccSteams.xlsx")

        print(f"Total de steams encontrados: {len(df_stems)}\n")

        # 2. Aplicar función hash a cada steam
        tabla_hash = {}

        for steam in df_stems["Token"]:
            valor_hash = hash(steam)
            tabla_hash[steam] = valor_hash

        # 3. Guardar archivo 5
        df_hash = pd.DataFrame(list(tabla_hash.items()), columns=["Steam", "Hash"])
        df_hash.to_excel("5ListDiccIndex.xlsx", index=False)

        print("✔ Archivo generado: 5ListDiccIndex.xlsx\n")

        # ============================================================
        # 4. PROCESAR BÚSQUEDA QUE PIDA EL USUARIO
        # ============================================================
        steam_buscar = request.form["steam_buscar"]
        steam_buscar = steam_buscar.lower()

        print(f"\nBuscando steam: {steam_buscar}")

        if steam_buscar in tabla_hash:
            print("🔎 Resultado: ENCONTRADO")
            print(f"   Hash: {tabla_hash[steam_buscar]}\n")
        else:
            print("❌ Resultado: NO ENCONTRADO\n")

        return "<h3>Indexación Hash realizada. Revisa la consola.</h3>"

    except Exception as e:
        print("\n❌ ERROR en indexación hash:", e, "\n")
        return "<h3>Error en la indexación. Revisa consola.</h3>"
    
    
@app.route("/procesar_arbol_b", methods=["POST"])
def procesar_arbol_b():
    palabra_query = request.form["palabra_b"].lower().strip()

    # Leer archivo Excel generado anteriormente
    import openpyxl
    wb = openpyxl.load_workbook("5ListDiccIndex.xlsx")
    ws = wb.active

    print("\n===== CREANDO ÁRBOL B =====")
    arbol = ArbolB(grado=4)

    steams = []

    for row in ws.iter_rows(min_row=2, values_only=True):
        steam = row[0]
        steams.append(steam)
        arbol.insertar(steam)

    print("\nÁRBOL B COMPLETO:")
    arbol.imprimir(arbol.raiz)

    print("\n===== BÚSQUEDA EN ÁRBOL B =====")
    print(f"Buscando: {palabra_query}")

    if arbol.buscar(arbol.raiz, palabra_query):
        print("Resultado: ENCONTRADO")
        return "<h3>Palabra ENCONTRADA en el árbol B.</h3>"
    else:
        print("Resultado: NO ENCONTRADA")
        return "<h3>Palabra NO encontrada en el árbol B.</h3>"

if __name__ == "__main__":
    app.run(debug=True)
